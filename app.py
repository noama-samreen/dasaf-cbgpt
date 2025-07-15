import streamlit as st
import json
import os
from cb_gpt_client import CbGptClient
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc import constants
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from io import BytesIO
import base64
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants
EDITED_RESPONSES_FILE = os.path.join(os.path.dirname(__file__), 'data', 'edited_responses.json')
RISKS_FILE = os.path.join(os.path.dirname(__file__), 'risks.json')

# Configure page settings
st.set_page_config(
    page_title="Blockchain Security Analysis Framework",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
    <style>
    .main { padding: 2rem; }
    .main-header {
        color: #1E3D59;
        padding: 1rem 0;
        border-bottom: 2px solid #f0f2f6;
        margin-bottom: 2rem;
    }
    .risk-section {
        background-color: #f8f9fa;
        padding: 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
    }
    </style>
""", unsafe_allow_html=True)

def get_download_link(buffer, filename, display_text):
    """Generate a download link for a file."""
    try:
        b64 = base64.b64encode(buffer.getvalue()).decode()
        return f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">{display_text}</a>'
    except Exception as e:
        logger.error(f"Error generating download link: {str(e)}")
        return None

def load_json_file(filepath):
    """Load and parse a JSON file."""
    try:
        with open(filepath, 'r') as f:
            return json.load(f)
    except FileNotFoundError:
        logger.error(f"File not found: {filepath}")
        # Return empty dict instead of failing
        return {}
    except json.JSONDecodeError as e:
        logger.error(f"Invalid JSON in file {filepath}: {str(e)}")
        return {}

def save_json_file(filepath, data):
    """Save data to a JSON file."""
    try:
        # Ensure directory exists
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        with open(filepath, 'w') as f:
            json.dump(data, f, indent=2)
    except Exception as e:
        logger.error(f"Error saving to {filepath}: {str(e)}")

def load_risks():
    """Load risks from JSON file."""
    return load_json_file(RISKS_FILE)

def load_edited_responses():
    """Load previously edited responses from file."""
    if not os.path.exists(EDITED_RESPONSES_FILE):
        # Create empty file if it doesn't exist
        save_json_file(EDITED_RESPONSES_FILE, {})
    return load_json_file(EDITED_RESPONSES_FILE)

def save_edited_response(risk_name, response):
    """Save an edited response to file and session state."""
    try:
        # Save to file
        responses = load_edited_responses()
        responses[risk_name] = response
        save_json_file(EDITED_RESPONSES_FILE, responses)
        
        # Update session state
        response_key = f"response_{risk_name}"
        st.session_state[response_key] = response
    except Exception as e:
        logger.error(f"Error saving edited response: {str(e)}")
        st.error("Failed to save changes. Please try again.")

def process_markdown(text, for_pdf=True):
    """Process markdown formatting for PDF or DOCX."""
    if not text:
        return text if for_pdf else (text, [], [])
    
    # Remove code block markers
    text = text.replace('```', '')
    
    if for_pdf:
        # Process inline formatting for PDF
        text = text.replace('`', '<font face="Courier" color="#333333">', 1)
        text = text.replace('`', '</font>', 1) if '`' in text else text
        
        # Process bold and italic
        text = text.replace('**', '<b>', 1)
        text = text.replace('**', '</b>', 1) if '**' in text else text
        text = text.replace('*', '<i>', 1)
        text = text.replace('*', '</i>', 1) if '*' in text else text
        
        # Process links
        import re
        text = re.sub(r'\[(.*?)\]\((.*?)\)', r'<link href="\2" color="blue"><u>\1</u></link>', text)
        
        return text
    else:
        # Process for DOCX
        formats = []
        links = []
        
        # Process inline code
        while '`' in text:
            start_idx = text.find('`')
            text = text.replace('`', '', 1)
            if '`' in text:
                end_idx = text.find('`')
                text = text.replace('`', '', 1)
                formats.append(('code', start_idx, end_idx))
        
        # Process bold
        while '**' in text:
            start_idx = text.find('**')
            text = text.replace('**', '', 1)
            if '**' in text:
                end_idx = text.find('**')
                text = text.replace('**', '', 1)
                formats.append(('bold', start_idx, end_idx))
        
        # Process italic
        while '*' in text:
            start_idx = text.find('*')
            text = text.replace('*', '', 1)
            if '*' in text:
                end_idx = text.find('*')
                text = text.replace('*', '', 1)
                formats.append(('italic', start_idx, end_idx))
        
        # Process links
        import re
        for match in re.finditer(r'\[(.*?)\]\((.*?)\)', text):
            links.append((match.group(1), match.group(2), match.start(), match.end()))
        
        # Replace link syntax with just the text
        text = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1', text)
        
        return text, formats, links

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

def add_formatted_paragraph(doc, text, style=None):
    """Add a paragraph with proper markdown formatting."""
    if not text.strip():
        return
    
    processed_text, formats, links = process_markdown(text, for_pdf=False)
    
    # Create paragraph
    paragraph = doc.add_paragraph()
    if style:
        paragraph.style = style
    
    # Handle code blocks
    if text.startswith('    ') or text.startswith('```'):
        run = paragraph.add_run(processed_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        return
    
    # Add formatted text
    current_pos = 0
    all_formats = [(start, end, 'format', fmt) for fmt, start, end in formats]
    all_formats.extend([(start, end, 'link', (text, url)) for text, url, start, end in links])
    all_formats.sort(key=lambda x: x[0])
    
    for start, end, type_, format_info in all_formats:
        # Add text before format
        if start > current_pos:
            paragraph.add_run(processed_text[current_pos:start])
        
        # Add formatted text
        if type_ == 'format':
            run = paragraph.add_run(processed_text[start:end])
            if format_info == 'bold':
                run.bold = True
            elif format_info == 'italic':
                run.italic = True
            elif format_info == 'code':
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        elif type_ == 'link':
            text, url = format_info
            add_hyperlink(paragraph, text, url)
        
        current_pos = end
    
    # Add remaining text
    if current_pos < len(processed_text):
        paragraph.add_run(processed_text[current_pos:])

def generate_docx(blockchain_name, blockchain_symbol, blockchain_website, critical_risks, non_critical_risks, edited_responses):
    """Generate a DOCX report of the security analysis."""
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Security Analysis Report: {blockchain_name} ({blockchain_symbol})', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add website
    if blockchain_website:
        paragraph = doc.add_paragraph()
        paragraph.add_run("Block Explorer: ")
        add_hyperlink(paragraph, blockchain_website, blockchain_website)
    doc.add_paragraph()  # Add spacing
    
    # Add critical risks section
    doc.add_heading('Critical Security Risks', 1)
    for risk in critical_risks:
        doc.add_heading(risk['name'], 2)
        content = edited_responses.get(risk['name'], 'No analysis available')
        if content and content.strip():
            # Split the analysis into paragraphs
            paragraphs = content.split('\n')
            for para in paragraphs:
                if para.strip():
                    add_formatted_paragraph(doc, para.strip())
        else:
            doc.add_paragraph("No analysis available")
        doc.add_paragraph()  # Add spacing
    
    # Add non-critical risks section
    doc.add_heading('Other Security Considerations', 1)
    for risk in non_critical_risks:
        doc.add_heading(risk['name'], 2)
        content = edited_responses.get(risk['name'], 'No analysis available')
        if content and content.strip():
            # Split the analysis into paragraphs
            paragraphs = content.split('\n')
            for para in paragraphs:
                if para.strip():
                    add_formatted_paragraph(doc, para.strip())
        else:
            doc.add_paragraph("No analysis available")
        doc.add_paragraph()  # Add spacing
    
    # Save to BytesIO
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

def generate_pdf(blockchain_name, blockchain_symbol, blockchain_website, critical_risks, non_critical_risks, edited_responses):
    """Generate a PDF report of the security analysis."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    
    # Define styles
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=1,  # Center alignment
        textColor=colors.HexColor('#1E3D59')
    )
    
    heading1_style = ParagraphStyle(
        'CustomHeading1',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=12,
        spaceBefore=24,
        textColor=colors.HexColor('#2E5575')
    )
    
    heading2_style = ParagraphStyle(
        'CustomHeading2',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=8,
        spaceBefore=16,
        textColor=colors.HexColor('#2E5575')
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=12,
        leading=14
    )
    
    code_style = ParagraphStyle(
        'CodeStyle',
        parent=styles['Code'],
        fontSize=10,
        fontName='Courier',
        spaceAfter=12,
        leading=14,
        backColor=colors.HexColor('#f5f5f5')
    )
    
    # Build the document content
    content = []
    
    # Add title
    content.append(Paragraph(f"Security Analysis Report:<br/>{blockchain_name} ({blockchain_symbol})", title_style))
    
    # Add website if provided
    if blockchain_website:
        content.append(Paragraph(f'Block Explorer: <link href="{blockchain_website}" color="blue"><u>{blockchain_website}</u></link>', normal_style))
    
    content.append(Spacer(1, 20))
    
    # Add critical risks section
    content.append(Paragraph("Critical Security Risks", heading1_style))
    
    for risk in critical_risks:
        content.append(Paragraph(risk['name'], heading2_style))
        analysis = edited_responses.get(risk['name'])
        if analysis and analysis.strip():
            # Split the analysis into paragraphs and process markdown in each
            paragraphs = analysis.split('\n')
            for para in paragraphs:
                if para.strip():
                    # Check if this is a code block (indented or between backticks)
                    if para.startswith('    ') or para.startswith('```'):
                        content.append(Paragraph(process_markdown(para.strip(), for_pdf=True), code_style))
                    else:
                        content.append(Paragraph(process_markdown(para.strip(), for_pdf=True), normal_style))
        else:
            content.append(Paragraph("No analysis available", normal_style))
        content.append(Spacer(1, 12))
    
    # Add non-critical risks section
    content.append(Paragraph("Other Security Considerations", heading1_style))
    
    for risk in non_critical_risks:
        content.append(Paragraph(risk['name'], heading2_style))
        analysis = edited_responses.get(risk['name'])
        if analysis and analysis.strip():
            # Split the analysis into paragraphs and process markdown in each
            paragraphs = analysis.split('\n')
            for para in paragraphs:
                if para.strip():
                    # Check if this is a code block (indented or between backticks)
                    if para.startswith('    ') or para.startswith('```'):
                        content.append(Paragraph(process_markdown(para.strip(), for_pdf=True), code_style))
                    else:
                        content.append(Paragraph(process_markdown(para.strip(), for_pdf=True), normal_style))
        else:
            content.append(Paragraph("No analysis available", normal_style))
        content.append(Spacer(1, 12))
    
    try:
        # Build the PDF
        doc.build(content)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"Error generating PDF: {str(e)}")
        return None

def display_risk_analysis(risk, response):
    """Display risk analysis with edit functionality."""
    st.markdown(f"### {risk['name']}")
    
    # Initialize session state keys
    edit_key = f"show_edit_{risk['name']}"
    response_key = f"response_{risk['name']}"
    
    # Initialize session state if not exists
    if edit_key not in st.session_state:
        st.session_state[edit_key] = False
    if response_key not in st.session_state:
        st.session_state[response_key] = response
    
    # Display analysis in expandable container
    with st.expander("View Analysis", expanded=True):
        # Show current analysis
        if st.session_state[response_key]:
            st.markdown(st.session_state[response_key])
        else:
            st.info("No analysis available yet.")
        
        # Edit button
        col1, col2, col3 = st.columns([1, 1, 3])
        
        with col1:
            if not st.session_state[edit_key]:
                if st.button("📝 Edit", key=f"edit_btn_{risk['name']}"):
                    st.session_state[edit_key] = True
                    st.experimental_rerun()
        
        # Show edit box when editing
        if st.session_state[edit_key]:
            edited_response = st.text_area(
                "Edit Analysis",
                value=st.session_state[response_key] if st.session_state[response_key] else "",
                key=f"edit_{risk['name']}",
                height=200
            )
            
            col1, col2, col3 = st.columns([1, 1, 3])
            
            with col1:
                if st.button("💾 Save", key=f"save_{risk['name']}"):
                    save_edited_response(risk['name'], edited_response)
                    st.session_state[response_key] = edited_response
                    st.session_state[edit_key] = False
                    st.success("✅ Changes saved!")
                    st.experimental_rerun()
            
            with col2:
                if st.button("❌ Cancel", key=f"cancel_{risk['name']}"):
                    st.session_state[edit_key] = False
                    st.experimental_rerun()
            
            with col3:
                if st.button("🔄 Regenerate", key=f"regen_{risk['name']}"):
                    client = CbGptClient()
                    new_response = client.analyze_blockchain_security(
                        st.session_state.blockchain_name,
                        risk['prompt'],
                        block_explorer_url=st.session_state.blockchain_website if hasattr(st.session_state, 'blockchain_website') else None
                    )
                    
                    if new_response:
                        save_edited_response(risk['name'], new_response)
                        st.session_state[response_key] = new_response
                        st.session_state[edit_key] = False
                        st.success("✅ Analysis regenerated!")
                        st.experimental_rerun()
                    else:
                        st.error("❌ Failed to regenerate analysis")

def main():
    """Main application function."""
    st.markdown("<h1 class='main-header'>Blockchain Security Analysis Framework</h1>", unsafe_allow_html=True)
    
    # Initialize session state for form data if not exists
    if 'form_submitted' not in st.session_state:
        st.session_state.form_submitted = False
    
    # Input form
    with st.form("blockchain_info"):
        col1, col2, col3 = st.columns(3)
        with col1:
            blockchain_name = st.text_input("Blockchain Name", 
                value=st.session_state.get('blockchain_name', ''))
        with col2:
            blockchain_symbol = st.text_input("Symbol", 
                value=st.session_state.get('blockchain_symbol', ''))
        with col3:
            blockchain_website = st.text_input(
                "Block Explorer URL", 
                value=st.session_state.get('blockchain_website', ''),
                help="Enter the blockchain's block explorer URL (e.g., https://etherscan.io) for validator distribution analysis"
            )
        
        submitted = st.form_submit_button("Analyze Security")
    
    # Handle form submission
    if submitted and blockchain_name and blockchain_symbol:
        # Store form data in session state
        st.session_state.blockchain_name = blockchain_name
        st.session_state.blockchain_symbol = blockchain_symbol
        st.session_state.blockchain_website = blockchain_website
        st.session_state.form_submitted = True
        
        # Clear existing responses
        if os.path.exists(EDITED_RESPONSES_FILE):
            os.remove(EDITED_RESPONSES_FILE)
        save_json_file(EDITED_RESPONSES_FILE, {})
        
        # Clear session state responses
        for key in list(st.session_state.keys()):
            if key.startswith('response_'):
                del st.session_state[key]
    
    # Show analysis if form was submitted (either now or previously)
    if st.session_state.form_submitted:
        # Display header
        st.header(f"Security Analysis for {st.session_state.blockchain_name} ({st.session_state.blockchain_symbol})")
        if st.session_state.blockchain_website:
            st.markdown(f"Block Explorer: [{st.session_state.blockchain_website}]({st.session_state.blockchain_website})")
        
        # Initialize client and load risks
        risks_data = load_risks()
        cb_gpt = CbGptClient()
        
        # Group risks by criticality
        critical_risks = [risk for risk in risks_data['risks'] if risk['is_critical']]
        non_critical_risks = [risk for risk in risks_data['risks'] if not risk['is_critical']]
        
        # Add reset button
        if st.button("⚠️ Start New Analysis"):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            if os.path.exists(EDITED_RESPONSES_FILE):
                os.remove(EDITED_RESPONSES_FILE)
            st.experimental_rerun()
            return
        
        # Load existing responses
        edited_responses = load_edited_responses()
        
        # Display risks
        st.markdown("## 🚨 Critical Security Risks")
        for risk in critical_risks:
            # Check if response exists in session state or edited responses
            response_key = f"response_{risk['name']}"
            if response_key not in st.session_state:
                # Check if response exists in edited responses file
                if risk['name'] in edited_responses:
                    st.session_state[response_key] = edited_responses[risk['name']]
                else:
                    # Generate new response
                    response = cb_gpt.analyze_blockchain_security(
                        st.session_state.blockchain_name,
                        risk['prompt'],
                        block_explorer_url=st.session_state.blockchain_website if st.session_state.blockchain_website else None
                    )
                    st.session_state[response_key] = response
                    # Save to edited responses
                    save_edited_response(risk['name'], response)
            
            display_risk_analysis(risk, st.session_state[response_key])
        
        st.markdown("## ⚠️ Other Security Considerations")
        for risk in non_critical_risks:
            # Check if response exists in session state or edited responses
            response_key = f"response_{risk['name']}"
            if response_key not in st.session_state:
                # Check if response exists in edited responses file
                if risk['name'] in edited_responses:
                    st.session_state[response_key] = edited_responses[risk['name']]
                else:
                    # Generate new response
                    response = cb_gpt.analyze_blockchain_security(
                        st.session_state.blockchain_name,
                        risk['prompt'],
                        block_explorer_url=st.session_state.blockchain_website if st.session_state.blockchain_website else None
                    )
                    st.session_state[response_key] = response
                    # Save to edited responses
                    save_edited_response(risk['name'], response)
            
            display_risk_analysis(risk, st.session_state[response_key])
        
        # Export functionality
        st.markdown("### Export Report")
        col1, col2 = st.columns(2)
        
        # Reload edited responses to ensure we have the latest version
        edited_responses = load_edited_responses()
        
        with col1:
            docx_buffer = generate_docx(
                st.session_state.blockchain_name,
                st.session_state.blockchain_symbol,
                st.session_state.blockchain_website,
                critical_risks,
                non_critical_risks,
                edited_responses
            )
            st.markdown(
                get_download_link(
                    docx_buffer,
                    f"{st.session_state.blockchain_name}_Security_Analysis.docx",
                    "📄 Download as DOCX"
                ),
                unsafe_allow_html=True
            )
        
        with col2:
            pdf_buffer = generate_pdf(
                st.session_state.blockchain_name,
                st.session_state.blockchain_symbol,
                st.session_state.blockchain_website,
                critical_risks,
                non_critical_risks,
                edited_responses
            )
            if pdf_buffer:
                st.markdown(
                    get_download_link(
                        pdf_buffer,
                        f"{st.session_state.blockchain_name}_Security_Analysis.pdf",
                        "📑 Download as PDF"
                    ),
                    unsafe_allow_html=True
                )
            else:
                st.error("Failed to generate PDF. Please check the debug information above.")

if __name__ == "__main__":
    main() 
